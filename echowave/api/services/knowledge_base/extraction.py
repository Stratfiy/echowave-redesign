"""Getting plain text, and where it came from, out of an uploaded file.

Scope is set by what the upload dialog accepts — PDF, DOCX, TXT, JSON, plus
the Markdown/CSV/HTML that arrive anyway as ``.txt`` cousins. Legacy ``.doc``
is accepted by the picker and cannot be read by any pure-Python library, so it
gets a named refusal rather than a generic one: telling somebody to re-save as
.docx is a fix they can carry out in thirty seconds.

Two properties matter more than fidelity here:

* **Page and heading provenance survives.** A chunk that knows it came from
  "Refund policy, page 4" retrieves better and cites better than one that does
  not, and the metadata columns already exist.
* **An empty extraction is a failure, not an empty success.** A scanned PDF
  with no text layer reads as zero characters. See
  :class:`~api.services.knowledge_base.errors.EmptyDocumentError`.

``docling`` is used when it is installed — it is what MPS used, and it reads
tables and multi-column layouts better than anything here. It is an optional
import rather than a requirement because it pulls a torch-sized dependency
tree, and the pure-Python path below is enough for the formats we accept.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import Any

from loguru import logger

from api.services.knowledge_base.errors import (
    DocumentExtractionError,
    EmptyDocumentError,
    UnsupportedDocumentTypeError,
)

#: Extensions handled by the pure-Python path, mapped to the reader used.
TEXT_EXTENSIONS = {".txt", ".text", ".log"}
MARKDOWN_EXTENSIONS = {".md", ".markdown"}
HTML_EXTENSIONS = {".html", ".htm"}
JSON_EXTENSIONS = {".json"}
CSV_EXTENSIONS = {".csv", ".tsv"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}

#: Accepted by the upload picker, unreadable without a converter binary.
LEGACY_WORD_EXTENSIONS = {".doc"}

SUPPORTED_EXTENSIONS = (
    TEXT_EXTENSIONS
    | MARKDOWN_EXTENSIONS
    | HTML_EXTENSIONS
    | JSON_EXTENSIONS
    | CSV_EXTENSIONS
    | PDF_EXTENSIONS
    | DOCX_EXTENSIONS
)

#: A page or paragraph shorter than this after stripping is whitespace noise —
#: a page number, a header rule — not content. Kept low deliberately: "Yes." is
#: a real answer in a policy document.
MIN_MEANINGFUL_CHARS = 2


@dataclass
class TextBlock:
    """One paragraph, list item or heading, with where it came from.

    ``heading_path`` is the stack of headings above this block, outermost
    first. It is what makes a chunk self-describing once it has been torn out
    of its document, and it is prepended to the embedded text.
    """

    text: str
    heading_path: tuple[str, ...] = ()
    page_number: int | None = None
    is_heading: bool = False


@dataclass
class ExtractedDocument:
    """Everything the chunker needs, and everything the row records."""

    blocks: list[TextBlock]
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks if block.text.strip())

    @property
    def character_count(self) -> int:
        return sum(len(block.text) for block in self.blocks)


class _HTMLTextExtractor(HTMLParser):
    """Tags out, structure kept.

    Headings become heading blocks so the same provenance rule applies to HTML
    as to everything else; ``script`` and ``style`` bodies are dropped, since
    embedding a stylesheet is worse than useless — it retrieves.
    """

    _SKIP = {"script", "style", "head", "meta", "link"}
    _HEADINGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
    _BREAKS = {"p", "div", "br", "li", "tr", "section", "article", "td", "th"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[tuple[str, int]] = []
        self._skip_depth = 0
        self._heading_level = 0
        self._buffer: list[str] = []

    def _flush(self, level: int = 0) -> None:
        text = " ".join("".join(self._buffer).split())
        self._buffer = []
        if text:
            self.parts.append((text, level))

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._SKIP:
            self._skip_depth += 1
            return
        if tag in self._HEADINGS:
            self._flush()
            self._heading_level = int(tag[1])
        elif tag in self._BREAKS:
            self._flush()

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if tag in self._HEADINGS:
            self._flush(self._heading_level)
            self._heading_level = 0
        elif tag in self._BREAKS:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            self._buffer.append(data)

    def close(self) -> None:  # type: ignore[override]
        super().close()
        self._flush(self._heading_level)


def _read_text_file(path: str) -> str:
    """Decode as UTF-8, then fall back rather than fail.

    A .txt uploaded from Windows is frequently cp1252, and refusing it over a
    smart quote would be an absurd reason to reject a policy document. Errors
    are replaced on the final attempt: a mangled character costs one bad token,
    a rejected file costs the whole feature.
    """
    with open(path, "rb") as handle:
        raw = handle.read()

    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")
#: A short line with no terminal punctuation, in a plain text file, is very
#: often a section title. Used only for .txt — Markdown and HTML say so
#: explicitly and PDFs carry font information we do not read.
_PLAIN_HEADING_MAX_CHARS = 80


def _blocks_from_markdown(text: str) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    heading_path: list[str] = []
    paragraph: list[str] = []
    in_code_fence = False

    def flush_paragraph() -> None:
        joined = " ".join(" ".join(paragraph).split())
        paragraph.clear()
        if len(joined) >= MIN_MEANINGFUL_CHARS:
            blocks.append(TextBlock(text=joined, heading_path=tuple(heading_path)))

    for line in text.splitlines():
        stripped = line.strip()

        if stripped.startswith("```"):
            # Inside a fence, a '#' is a comment, not a heading, and blank
            # lines are part of the listing. Keep the block whole.
            in_code_fence = not in_code_fence
            if not in_code_fence:
                flush_paragraph()
            continue

        if in_code_fence:
            paragraph.append(line)
            continue

        if not stripped:
            flush_paragraph()
            continue

        match = _MARKDOWN_HEADING.match(stripped)
        if match:
            flush_paragraph()
            level = len(match.group(1))
            title = match.group(2)
            del heading_path[level - 1 :]
            heading_path.append(title)
            blocks.append(
                TextBlock(
                    text=title, heading_path=tuple(heading_path[:-1]), is_heading=True
                )
            )
            continue

        paragraph.append(stripped)

    flush_paragraph()
    return blocks


def _blocks_from_plain_text(text: str) -> list[TextBlock]:
    """Paragraphs split on blank lines, with a light heading heuristic."""
    blocks: list[TextBlock] = []
    heading_path: list[str] = []

    for raw_paragraph in re.split(r"\n\s*\n", text):
        paragraph = " ".join(raw_paragraph.split())
        if len(paragraph) < MIN_MEANINGFUL_CHARS:
            continue

        looks_like_heading = (
            "\n" not in raw_paragraph.strip()
            and len(paragraph) <= _PLAIN_HEADING_MAX_CHARS
            and not paragraph.endswith((".", "!", "?", ",", ";", ":"))
        )
        if looks_like_heading:
            heading_path = [paragraph]
            blocks.append(TextBlock(text=paragraph, is_heading=True))
        else:
            blocks.append(TextBlock(text=paragraph, heading_path=tuple(heading_path)))

    return blocks


def _blocks_from_html(text: str) -> list[TextBlock]:
    parser = _HTMLTextExtractor()
    parser.feed(text)
    parser.close()

    blocks: list[TextBlock] = []
    heading_path: list[str] = []
    for part, level in parser.parts:
        if len(part) < MIN_MEANINGFUL_CHARS:
            continue
        if level:
            del heading_path[level - 1 :]
            heading_path.append(part)
            blocks.append(
                TextBlock(
                    text=part, heading_path=tuple(heading_path[:-1]), is_heading=True
                )
            )
        else:
            blocks.append(TextBlock(text=part, heading_path=tuple(heading_path)))
    return blocks


def _blocks_from_json(text: str) -> list[TextBlock]:
    """Flatten to ``key.path: value`` lines.

    A JSON export embedded as raw source retrieves badly — the braces and
    quoting dominate the vector. One line per leaf, carrying its path, is both
    readable and searchable. Invalid JSON falls back to plain text rather than
    failing: a .json that is really a JSONL log is still worth ingesting.
    """
    try:
        parsed = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        logger.debug("JSON document did not parse; ingesting as plain text")
        return _blocks_from_plain_text(text)

    lines: list[str] = []

    def walk(node: Any, path: str) -> None:
        if isinstance(node, dict):
            for key, value in node.items():
                walk(value, f"{path}.{key}" if path else str(key))
        elif isinstance(node, list):
            for index, value in enumerate(node):
                walk(value, f"{path}[{index}]")
        else:
            rendered = "" if node is None else str(node)
            if rendered.strip():
                lines.append(f"{path}: {rendered}" if path else rendered)

    walk(parsed, "")
    return [TextBlock(text=line) for line in lines]


def _blocks_from_csv(text: str, *, delimiter: str) -> list[TextBlock]:
    """One block per row, with the header repeated as ``column: value``.

    A bare row of values loses its meaning the moment it leaves the table.
    Repeating the header costs tokens and buys a chunk that answers "what is
    the refund window for plan B" instead of retrieving the digit 14.
    """
    import csv
    import io

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = [row for row in reader if any(cell.strip() for cell in row)]
    if not rows:
        return []

    header, *body = rows
    if not body:
        return [TextBlock(text=", ".join(cell.strip() for cell in header))]

    blocks = []
    for row in body:
        pairs = [
            f"{(header[i] if i < len(header) else f'column {i + 1}').strip()}: {cell.strip()}"
            for i, cell in enumerate(row)
            if cell.strip()
        ]
        if pairs:
            blocks.append(TextBlock(text="; ".join(pairs)))
    return blocks


def _blocks_from_pdf(path: str, metadata: dict[str, Any]) -> list[TextBlock]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentExtractionError(
            "pypdf is not installed in this image",
            user_message=(
                "PDF support is not available on this deployment. Upload the "
                "document as .docx or .txt, or contact support."
            ),
        ) from exc

    try:
        reader = PdfReader(path)
        if reader.is_encrypted:
            # An empty user password is the common "protected against editing"
            # case and decrypts fine. A real password does not, and that is a
            # refusal the customer can act on.
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentExtractionError(
                    f"PDF is password protected: {exc}",
                    user_message=(
                        "This PDF is password protected, so we cannot read it. "
                        "Remove the password and upload it again."
                    ),
                ) from exc
        pages = reader.pages
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError(
            f"Could not open PDF: {exc}",
            user_message=(
                "We could not open this PDF — it may be corrupt or "
                "incompletely uploaded. Try uploading it again."
            ),
        ) from exc

    blocks: list[TextBlock] = []
    pages_with_text = 0
    for page_index, page in enumerate(pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            # One unreadable page must not lose the other ninety-nine.
            logger.warning(f"PDF page {page_index} could not be read: {exc}")
            continue

        if len(page_text.strip()) >= MIN_MEANINGFUL_CHARS:
            pages_with_text += 1

        for raw_paragraph in re.split(r"\n\s*\n", page_text):
            paragraph = " ".join(raw_paragraph.split())
            if len(paragraph) >= MIN_MEANINGFUL_CHARS:
                blocks.append(TextBlock(text=paragraph, page_number=page_index))

    metadata["page_count"] = len(pages)
    metadata["pages_with_text"] = pages_with_text
    #: The signal that separates "scanned image" from "document we mishandled".
    metadata["has_text_layer"] = pages_with_text > 0
    return blocks


def _blocks_from_docx(path: str, metadata: dict[str, Any]) -> list[TextBlock]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise DocumentExtractionError(
            "python-docx is not installed in this image",
            user_message=(
                "Word support is not available on this deployment. Upload the "
                "document as PDF or .txt, or contact support."
            ),
        ) from exc

    try:
        document = docx.Document(path)
    except Exception as exc:
        raise DocumentExtractionError(
            f"Could not open DOCX: {exc}",
            user_message=(
                "We could not open this Word file — it may be corrupt, or "
                "saved in the older .doc format. Re-save it as .docx and "
                "upload it again."
            ),
        ) from exc

    blocks: list[TextBlock] = []
    heading_path: list[str] = []

    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if len(text) < MIN_MEANINGFUL_CHARS:
            continue

        style_name = (paragraph.style.name if paragraph.style else "") or ""
        heading_match = re.match(r"Heading (\d)", style_name)
        if heading_match or style_name == "Title":
            level = 1 if style_name == "Title" else int(heading_match.group(1))
            del heading_path[level - 1 :]
            heading_path.append(text)
            blocks.append(
                TextBlock(
                    text=text, heading_path=tuple(heading_path[:-1]), is_heading=True
                )
            )
        else:
            blocks.append(TextBlock(text=text, heading_path=tuple(heading_path)))

    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            [" ".join(cell.text.split()) for cell in row.cells] for row in table.rows
        ]
        rows = [row for row in rows if any(cell for cell in rows[0]) and any(row)]
        if not rows:
            continue
        header, *body = rows
        if not body:
            blocks.append(
                TextBlock(
                    text=f"Table {table_index}: " + ", ".join(c for c in header if c),
                    heading_path=tuple(heading_path),
                )
            )
            continue
        for row in body:
            pairs = [
                f"{(header[i] if i < len(header) else f'column {i + 1}')}: {cell}"
                for i, cell in enumerate(row)
                if cell
            ]
            if pairs:
                blocks.append(
                    TextBlock(text="; ".join(pairs), heading_path=tuple(heading_path))
                )

    metadata["paragraph_count"] = len(document.paragraphs)
    metadata["table_count"] = len(document.tables)
    return blocks


def _docling_available() -> bool:
    """Whether the higher-fidelity converter is in this image.

    Checked by import rather than by a setting, because a setting that claims
    docling is present on an image that lacks it fails at the worst moment —
    mid-ingestion, on a customer's file.
    """
    try:
        import docling  # noqa: F401
    except Exception:
        return False
    return True


def _blocks_from_docling(path: str, metadata: dict[str, Any]) -> list[TextBlock] | None:
    """Convert via docling, or return ``None`` to fall through.

    Returns rather than raises on failure: docling not coping with a file is
    not a reason to reject a file pypdf reads fine.
    """
    try:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(path)
        markdown = result.document.export_to_markdown()
    except Exception as exc:
        logger.warning(f"docling conversion failed, falling back: {exc}")
        return None

    if not markdown or not markdown.strip():
        return None

    metadata["extractor"] = "docling"
    return _blocks_from_markdown(markdown)


def _blocks_from_ocr(path: str, metadata: dict[str, Any]) -> list[TextBlock]:
    """Recognised text from a scanned PDF, one block per page.

    Empty when OCR is unavailable on this deployment or recognised nothing, in
    which case the caller raises the same "this looks like a scan" message it
    always did. ``extractor`` is overwritten so the provenance on every chunk
    says the text was read by machine rather than lifted from a text layer —
    OCR is approximate, and a quote the agent reads out should be traceable to
    that.
    """
    from api.services.knowledge_base import ocr

    pages = ocr.pages_from_pdf(path)
    if not pages:
        return []

    metadata["extractor"] = "tesseract"
    metadata["ocr_pages"] = len(pages)
    return [
        TextBlock(text=text, page_number=number)
        for number, text in enumerate(pages, start=1)
    ]


def extract_document(
    file_path: str,
    filename: str,
    content_type: str | None = None,
) -> ExtractedDocument:
    """Read ``file_path`` into blocks of text with their provenance.

    ``filename`` decides the reader — the browser-supplied content type is
    advisory and frequently ``application/octet-stream`` from a presigned PUT,
    so it is recorded and not trusted.

    Raises :class:`UnsupportedDocumentTypeError` for a format we cannot read,
    :class:`DocumentExtractionError` for one we should be able to read and
    could not, and :class:`EmptyDocumentError` when the file yields no text.
    """
    extension = os.path.splitext(filename)[1].lower()
    metadata: dict[str, Any] = {
        "backend": "local",
        "extractor": "builtin",
        "source_filename": filename,
        "source_extension": extension,
        "declared_content_type": content_type,
    }

    if extension in LEGACY_WORD_EXTENSIONS:
        raise UnsupportedDocumentTypeError(
            f"Legacy .doc is not readable in-process: {filename}",
            user_message=(
                "The older .doc format cannot be read. Open the file in Word "
                "and save it as .docx, or export it as PDF, then upload again."
            ),
        )

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedDocumentTypeError(
            f"Unsupported extension {extension!r} for {filename}",
            user_message=(
                f"We cannot read {extension or 'this file type'} documents. "
                f"Supported formats are: {supported}."
            ),
        )

    blocks: list[TextBlock] | None = None

    if extension in PDF_EXTENSIONS or extension in DOCX_EXTENSIONS:
        if _docling_available():
            blocks = _blocks_from_docling(file_path, metadata)

    if blocks is None:
        if extension in PDF_EXTENSIONS:
            metadata["extractor"] = "pypdf"
            blocks = _blocks_from_pdf(file_path, metadata)
        elif extension in DOCX_EXTENSIONS:
            metadata["extractor"] = "python-docx"
            blocks = _blocks_from_docx(file_path, metadata)
        elif extension in JSON_EXTENSIONS:
            blocks = _blocks_from_json(_read_text_file(file_path))
        elif extension in CSV_EXTENSIONS:
            delimiter = "\t" if extension == ".tsv" else ","
            blocks = _blocks_from_csv(_read_text_file(file_path), delimiter=delimiter)
        elif extension in HTML_EXTENSIONS:
            blocks = _blocks_from_html(_read_text_file(file_path))
        elif extension in MARKDOWN_EXTENSIONS:
            blocks = _blocks_from_markdown(_read_text_file(file_path))
        else:
            blocks = _blocks_from_plain_text(_read_text_file(file_path))

    blocks = [block for block in blocks if block.text.strip()]
    metadata["block_count"] = len(blocks)
    metadata["character_count"] = sum(len(block.text) for block in blocks)

    # A PDF with no text layer is a scan — pictures of words. Recognising them
    # is slow and approximate, so it happens only here, after the readers that
    # would have done better have already come back with nothing.
    if (
        extension in PDF_EXTENSIONS
        and metadata["character_count"] < MIN_MEANINGFUL_CHARS
    ):
        recognised = _blocks_from_ocr(file_path, metadata)
        if recognised:
            blocks = recognised
            metadata["block_count"] = len(blocks)
            metadata["character_count"] = sum(len(block.text) for block in blocks)

    if not blocks or metadata["character_count"] < MIN_MEANINGFUL_CHARS:
        # The scanned-PDF case gets its own sentence. Telling somebody their
        # PDF "contains no text" when they can plainly see text on the page is
        # the kind of accurate-but-useless message that generates a ticket.
        if extension in PDF_EXTENSIONS and not metadata.get("has_text_layer", False):
            raise EmptyDocumentError(
                f"PDF has no text layer: {filename}",
                user_message=(
                    "This PDF has no selectable text — it looks like a scan or "
                    "a set of page images. Run it through OCR (most scanners "
                    "and PDF tools offer 'searchable PDF'), then upload it "
                    "again."
                ),
            )
        raise EmptyDocumentError(
            f"No extractable text in {filename}",
            user_message=(
                "This file contains no text we can read, so there would be "
                "nothing for the agent to answer from. Check that the document "
                "is not empty and upload it again."
            ),
        )

    return ExtractedDocument(blocks=blocks, metadata=metadata)
